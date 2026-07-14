"""
Document Router (v3.1) — Multi-format text extraction.

The whole pipeline after the Fetcher works on PLAIN TEXT. So supporting a
new file type only requires a new parser here — chunker, RAPTOR, retriever,
synthesizer, critic all stay untouched.

Design decision: routing sniffs MAGIC BYTES before trusting the suffix —
modern arXiv PDF URLs have no .pdf extension, and a URL's suffix says nothing
about what the server actually returned.

Supported: PDF, DOCX, HTML, TXT/MD, CSV.
"""
import logging
from pathlib import Path
from config import settings

logger = logging.getLogger(__name__)


def _detect_kind(source: str, raw_bytes: bytes | None, suffix: str) -> str:
    if raw_bytes:
        if raw_bytes[:5] == b"%PDF-":
            return "pdf"
        if raw_bytes[:4] == b"PK\x03\x04" and suffix == "docx":
            return "docx"
        head = raw_bytes[:512].lstrip().lower()
        if head.startswith((b"<!doctype html", b"<html")):
            return "html"
    if suffix == "pdf":
        return "pdf"
    if suffix == "docx":
        return "docx"
    if suffix in ("html", "htm"):
        return "html"
    if suffix == "csv":
        return "csv"
    if suffix in ("txt", "md", ""):
        # extension-less URL with non-HTML, non-PDF body -> treat as text;
        # extension-less *URL with no bytes yet* is fetched as HTML.
        if raw_bytes is None and source.startswith("http"):
            return "html"
        return "text"
    return "text"


def extract_text(source: str, raw_bytes: bytes = None) -> str:
    """
    Route a file path / URL / raw bytes to the right parser.
    Returns markdown-ish plain text, truncated to max_paper_chars.
    Returns "" on failure (graceful degradation — caller skips the doc).
    """
    suffix = Path(source.split("?")[0]).suffix.lower().lstrip(".")
    if suffix and suffix not in settings.supported_formats and suffix not in ("htm",):
        logger.warning(f"Unsupported format '{suffix}' — attempting content sniffing")

    kind = _detect_kind(source, raw_bytes, suffix)
    try:
        if kind == "pdf":
            text = _parse_pdf(source, raw_bytes)
        elif kind == "docx":
            text = _parse_docx(source, raw_bytes)
        elif kind == "html":
            text = _parse_html(source, raw_bytes)
        elif kind == "csv":
            text = _parse_csv(source, raw_bytes)
        else:
            text = _parse_text(source, raw_bytes)
    except Exception as e:
        logger.error(f"Extraction failed for {source} ({kind}): {e}")
        return ""

    return (text or "")[: settings.max_paper_chars]


def _parse_pdf(source, raw_bytes):
    import pymupdf
    import pymupdf4llm
    if raw_bytes:
        # pymupdf4llm takes a path or a Document — never raw bytes directly.
        doc = pymupdf.open(stream=raw_bytes, filetype="pdf")
        return pymupdf4llm.to_markdown(doc)
    return pymupdf4llm.to_markdown(source)


def _parse_docx(source, raw_bytes):
    from docx import Document
    import io
    doc = Document(io.BytesIO(raw_bytes)) if raw_bytes else Document(source)
    parts = []
    for p in doc.paragraphs:
        style = (p.style.name or "").lower()
        if "heading" in style:
            level = "".join(ch for ch in style if ch.isdigit()) or "2"
            parts.append("#" * int(level) + " " + p.text)
        else:
            parts.append(p.text)
    return "\n\n".join(t for t in parts if t.strip())


def _parse_html(source, raw_bytes):
    import trafilatura
    if raw_bytes:
        return trafilatura.extract(raw_bytes.decode("utf-8", errors="ignore")) or ""
    downloaded = trafilatura.fetch_url(source)
    return trafilatura.extract(downloaded) or ""


def _parse_csv(source, raw_bytes):
    import pandas as pd
    import io
    df = pd.read_csv(io.BytesIO(raw_bytes)) if raw_bytes else pd.read_csv(source)
    return df.head(200).to_markdown(index=False)


def _parse_text(source, raw_bytes):
    if raw_bytes:
        return raw_bytes.decode("utf-8", errors="ignore")
    return Path(source).read_text(encoding="utf-8", errors="ignore")
